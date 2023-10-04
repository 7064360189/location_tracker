import requests
from bs4 import BeautifulSoup as BS
import docx



r=requests.get('https://www.javatpoint.com/python-interview-questions')
soup = BS(r.content, 'html.parser')
s=soup.find('div', class_='onlycontent')
content = s.find_all('h3')

# z=zip(content,ans)
# a=print(list(z))
# print(a)

d=docx.Document()
# d1=docx.Document()
# for i,a in z:
#     d1.add_pragraph(i.text,a.text)

# d=docx.Document()   
# for i in zip(content,ans):
#     for a in i:
#         # print(a.text)
#         d.add_paragraph(a.text)
#         d.save("C:\\Users\\user\\New folder\\python.docx")

    




for i in content:
    d.add_paragraph(i.text)
    d.save("C:\\Users\\user\\New folder\\python.docx")


# for a in ans:
#     d1.add_paragraph(a.text)
#     d1.save("c:\\Users\\user\\New folder\\python.docx")

# z=zip(content,ans)        
# print(z)
# for i,x in z:
#      d.add_paragraph(i.text,x.text)
#      d.save("C:\\Users\\user\\New folder\\python.docx")

